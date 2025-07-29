# super_agent/course_designer_views.py
import json
import google.generativeai as genai
from django.http import JsonResponse, StreamingHttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.conf import settings
from typing import Dict, List, Any, Optional, Generator
import time
from datetime import datetime
import logging
import tempfile
import os
import re

# 파일 텍스트 추출을 위한 라이브러리
import PyPDF2
import docx
import pandas as pd
from io import BytesIO

# 로거 설정
logger = logging.getLogger(__name__)

# Gemini API 설정
genai.configure(api_key=settings.GEMINI_API_KEY)


class FileTextExtractor:
    """파일에서 텍스트를 추출하는 클래스"""
    
    @staticmethod
    def extract_text(file) -> str:
        """파일 타입에 따라 텍스트 추출"""
        try:
            file_type = file.content_type
            file_name = file.name.lower()
            
            logger.info(f"[FileTextExtractor] 파일 텍스트 추출 시작: {file_name} ({file_type})")
            
            if file_type == 'application/pdf' or file_name.endswith('.pdf'):
                return FileTextExtractor.extract_pdf_text(file)
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                               'application/msword'] or file_name.endswith(('.docx', '.doc')):
                return FileTextExtractor.extract_word_text(file)
            elif file_type in ['application/vnd.ms-excel', 
                               'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                               'text/csv'] or file_name.endswith(('.xlsx', '.xls', '.csv')):
                return FileTextExtractor.extract_spreadsheet_text(file)
            elif file_type.startswith('text/') or file_name.endswith('.txt'):
                return file.read().decode('utf-8', errors='ignore')
            else:
                return f"[파일 타입 {file_type}는 텍스트 추출이 지원되지 않습니다]"
                
        except Exception as e:
            logger.error(f"[FileTextExtractor] 텍스트 추출 오류: {str(e)}")
            return f"[텍스트 추출 실패: {str(e)}]"
    
@staticmethod
def extract_pdf_text(file) -> str:
    """PDF에서 텍스트 추출 - 표 데이터 처리 개선"""
    try:
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        
        # 표 형식 데이터를 더 잘 처리하기 위해 pdfplumber 사용
        import pdfplumber
        
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                # 표 추출
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        # 표를 구조화된 텍스트로 변환
                        for row in table:
                            if row:  # None이 아닌 행만 처리
                                cleaned_row = [str(cell).strip() if cell else '' for cell in row]
                                text.append(' | '.join(cleaned_row))
                else:
                    # 표가 없으면 일반 텍스트 추출
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        
        return "\n".join(text)
        
    except Exception as e:
        # pdfplumber가 실패하면 PyPDF2로 폴백
        logger.warning(f"pdfplumber 실패, PyPDF2로 폴백: {str(e)}")
        return FileTextExtractor.extract_pdf_text_pypdf2(file)

    @staticmethod
    def extract_pdf_text_07211809(file) -> str:
        """PDF에서 텍스트 추출"""
        try:
            file.seek(0)  # 파일 포인터를 처음으로
            pdf_reader = PyPDF2.PdfReader(file)
            text = []
            
            total_pages = len(pdf_reader.pages)
            logger.info(f"[FileTextExtractor] PDF 페이지 수: {total_pages}")
            
            for page_num in range(min(total_pages, 50)):  # 최대 50페이지까지
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(f"--- 페이지 {page_num + 1} ---\n{page_text}")
            
            if total_pages > 50:
                text.append(f"\n[총 {total_pages}페이지 중 처음 50페이지만 추출했습니다]")
            
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"[FileTextExtractor] PDF 텍스트 추출 오류: {str(e)}")
            return f"[PDF 텍스트 추출 오류: {str(e)}]"
    
    @staticmethod
    def extract_word_text(file) -> str:
        """Word 문서에서 텍스트 추출"""
        try:
            file.seek(0)
            doc = docx.Document(file)
            paragraphs = []
            
            # 단락 추출
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 표 내용도 추출
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    paragraphs.append("\n[표]\n" + "\n".join(table_text))
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"[FileTextExtractor] Word 텍스트 추출 오류: {str(e)}")
            return f"[Word 텍스트 추출 오류: {str(e)}]"
    
    @staticmethod
    def extract_spreadsheet_text(file) -> str:
        """엑셀/CSV에서 텍스트 추출"""
        try:
            file.seek(0)
            
            if file.name.endswith('.csv'):
                df = pd.read_csv(file, encoding='utf-8-sig', on_bad_lines='skip')
            else:
                df = pd.read_excel(file)
            
            # 데이터프레임을 구조화된 텍스트로 변환
            text_parts = [f"[스프레드시트 정보]"]
            text_parts.append(f"총 {len(df)}행, {len(df.columns)}열의 데이터")
            text_parts.append(f"컬럼: {', '.join(df.columns)}")
            
            # 샘플 데이터 (최대 100행)
            sample_size = min(len(df), 100)
            if sample_size > 0:
                text_parts.append(f"\n[샘플 데이터 - 처음 {sample_size}행]")
                sample_df = df.head(sample_size)
                text_parts.append(sample_df.to_string(max_rows=100, max_cols=10))
            
            if len(df) > 100:
                text_parts.append(f"\n[전체 {len(df)}행 중 처음 100행만 표시]")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"[FileTextExtractor] 스프레드시트 텍스트 추출 오류: {str(e)}")
            return f"[스프레드시트 텍스트 추출 오류: {str(e)}]"


class CourseDesigner:
    def __init__(self, model_type: str = 'flash'):
        """
        CourseDesigner 초기화
        
        Args:
            model_type: 'flash' 또는 'pro' 선택
        """
        logger.info(f"[CourseDesigner] 초기화 시작 - 모델 타입: {model_type}")
        
        # 모델 선택
        try:
            if model_type == 'pro':
                self.model = genai.GenerativeModel('gemini-2.5-pro')
            else:
                self.model = genai.GenerativeModel('gemini-2.5-flash')
            
            logger.info(f"[CourseDesigner] 모델 초기화 성공")
        except Exception as e:
            logger.error(f"[CourseDesigner] 모델 초기화 실패: {str(e)}")
            raise
        
        self.model_type = model_type
        
        # 교육 단계별 매핑
        self.education_level_map = {
            'elementary': '초등학교',
            'middle': '중학교',
            'high': '고등학교'
        }
        
        # 난이도 매핑
        self.difficulty_map = {
            '상': '심화/고급',
            '중': '표준/중급',
            '하': '기초/초급'
        }

    def process_files_to_text(self, files) -> str:  
        """파일들을 텍스트로 변환하여 하나의 문자열로 결합"""  
        extracted_texts = []
        
        for idx, file in enumerate(files):
            logger.info(f"[process_files] 파일 {idx+1}/{len(files)} 텍스트 추출 중: {file.name}")
            
            try:
                # 파일 텍스트 추출
                text = FileTextExtractor.extract_text(file)
                
                if text and not text.startswith("["):
                    extracted_texts.append(f"=== 파일 {idx+1}: {file.name} ===\n{text}")
                    logger.info(f"[process_files] 파일 {file.name} 텍스트 추출 성공 - {len(text)} 글자")
                else:
                    logger.warning(f"[process_files] 파일 {file.name} 텍스트 추출 실패: {text}")
                    
            except Exception as e:
                logger.error(f"[process_files] 파일 {file.name} 처리 오류: {str(e)}")
        
        if not extracted_texts:
            logger.warning("[process_files] 추출된 텍스트가 없습니다")
            return ""
        
        combined_text = "\n\n".join(extracted_texts)
        original_length = len(combined_text)
        
        # 전체 코스 생성 모드인 경우 텍스트 길이 제한을 늘림
        max_length = 100000  # 100KB로 증가
        if original_length > max_length:
            combined_text = combined_text[:max_length] + f"\n\n[전체 {original_length}자 중 {max_length}자까지만 포함]"
            logger.info(f"[process_files] 텍스트 길이 제한: {original_length} → {max_length}")
        
        logger.info(f"[process_files] 전체 추출된 텍스트 길이: {len(combined_text)} 글자")
        return combined_text
    
    def process_files_to_text_07211822(self, files) -> str:
        """파일들을 텍스트로 변환하여 하나의 문자열로 결합"""
        extracted_texts = []
        
        for idx, file in enumerate(files):
            logger.info(f"[process_files] 파일 {idx+1}/{len(files)} 텍스트 추출 중: {file.name}")
            
            try:
                # 파일 텍스트 추출
                text = FileTextExtractor.extract_text(file)
                
                if text and not text.startswith("["):  # 에러 메시지가 아닌 경우
                    extracted_texts.append(f"=== 파일 {idx+1}: {file.name} ===\n{text}")
                    logger.info(f"[process_files] 파일 {file.name} 텍스트 추출 성공 - {len(text)} 글자")
                else:
                    logger.warning(f"[process_files] 파일 {file.name} 텍스트 추출 실패: {text}")
                    
            except Exception as e:
                logger.error(f"[process_files] 파일 {file.name} 처리 오류: {str(e)}")
        
        if not extracted_texts:
            logger.warning("[process_files] 추출된 텍스트가 없습니다")
            return ""
        
        combined_text = "\n\n".join(extracted_texts)
        original_length = len(combined_text)
        
        # 텍스트가 너무 길면 요약 또는 잘라내기
        max_length = 50000  # 50KB
        if original_length > max_length:
            combined_text = combined_text[:max_length] + f"\n\n[전체 {original_length}자 중 {max_length}자까지만 포함]"
            logger.info(f"[process_files] 텍스트 길이 제한: {original_length} → {max_length}")
        
        logger.info(f"[process_files] 전체 추출된 텍스트 길이: {len(combined_text)} 글자")
        return combined_text
    
    def create_concise_course_prompt(self, 
                                   user_prompt: str,
                                   extracted_text: str = "",
                                   education_level: str = "middle",
                                   difficulty: str = "중",
                                   include_assessment: bool = True,
                                   include_activities: bool = True,
                                   include_resources: bool = False) -> str:
        """더 간결한 프롬프트 생성 (토큰 절약)"""
        
        edu_level_kr = self.education_level_map.get(education_level, '중학교')
        diff_kr = self.difficulty_map.get(difficulty, '표준/중급')
        
        # 파일 텍스트는 요약해서 포함
        file_context = ""
        if extracted_text:
            # 텍스트가 길면 처음 부분만 사용
            max_context = 5000  # 5000자로 제한
            if len(extracted_text) > max_context:
                extracted_text = extracted_text[:max_context] + "\n[이하 생략]"
            file_context = f"\n참고자료:\n{extracted_text}\n"
        
        prompt = f"""
{edu_level_kr} {user_prompt} 코스를 JSON으로 생성하세요.
{file_context}
규칙:
- 대단원 2-3개, 소단원 각 2-3개, 차시 각 3-4개
- 각 차시 45분
- 코드블록 없이 순수 JSON만

형식:
{{
    "course": {{
        "subject_name": "코스명",
        "target": "{edu_level_kr} X학년",
        "description": "설명"
    }},
    "chapters": [
        {{
            "chapter_title": "대단원명",
            "chapter_order": 1,
            "description": "설명",
            "subchapters": [
                {{
                    "sub_chapter_title": "소단원명",
                    "sub_chapter_order": 1,
                    "description": "설명",
                    "chasis": [
                        {{
                            "chasi_title": "차시명",
                            "chasi_order": 1,
                            "description": "설명",
                            "learning_objectives": "목표",
                            "duration_minutes": 45
                        }}
                    ]
                }}
            ]
        }}
    ]
}}"""
        
        return prompt
    
    def create_full_course_prompt(self, 
                                user_prompt: str,
                                extracted_text: str = "",
                                education_level: str = "middle",
                                difficulty: str = "중",
                                **kwargs) -> str:
        """PDF의 전체 과정을 생성하는 프롬프트"""
        
        edu_level_kr = self.education_level_map.get(education_level, '중학교')
        
        prompt = f"""
{edu_level_kr} {user_prompt} 코스를 JSON으로 생성하세요.

참고자료에서 추출한 전체 차시 목록:
{extracted_text}

중요: 참고자료에 있는 모든 단원과 차시를 빠짐없이 포함시켜주세요.
표에 나온 순서와 구조를 정확히 따라서 생성해주세요.

형식:
{{
    "course": {{
        "subject_name": "국어 5-2",
        "target": "{edu_level_kr} 5학년",
        "description": "설명"
    }},
    "chapters": [
        {{
            "chapter_title": "단원명 (참고자료와 동일하게)",
            "chapter_order": 1,
            "description": "단원 설명",
            "subchapters": [
                {{
                    "sub_chapter_title": "소단원명",
                    "sub_chapter_order": 1,
                    "description": "소단원 설명",
                    "chasis": [
                        {{
                            "chasi_title": "차시명 (참고자료와 동일하게)",
                            "chasi_order": 1,
                            "description": "차시 설명",
                            "learning_objectives": "학습 목표",
                            "duration_minutes": 45
                        }}
                    ]
                }}
            ]
        }}
    ]
}}

규칙:
1. 참고자료에 있는 모든 단원을 포함 (국어(가), 국어(나) 구분 포함)
2. 차시 번호와 제목은 참고자료와 완전히 동일하게
3. 누락된 차시가 없도록 주의
4. 코드블록 없이 순수 JSON만 출력
"""
        
        return prompt

    def clean_json_response(self, text: str) -> str:
        """JSON 응답을 정리하고 수정하는 함수"""
        logger.info("[clean_json] JSON 응답 정리 시작")
        
        # 코드 블록 제거 (```json ... ```)
        text = re.sub(r'```json\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'```\s*', '', text)
        
        # 앞뒤 공백 제거
        text = text.strip()
        
        # JSON 시작과 끝 찾기
        start_idx = text.find('{')
        end_idx = text.rfind('}')
        
        if start_idx == -1 or end_idx == -1:
            logger.error("[clean_json] JSON 시작 또는 끝을 찾을 수 없음")
            return text
        
        json_text = text[start_idx:end_idx + 1]
        
        # 일반적인 JSON 오류 수정
        # 1. 마지막 쉼표 제거
        json_text = re.sub(r',\s*}', '}', json_text)
        json_text = re.sub(r',\s*]', ']', json_text)
        
        # 2. 이미 이스케이프된 줄바꿈은 유지, 그렇지 않은 것만 변환
        json_text = re.sub(r'(?<!\\)\\n', '\\\\n', json_text)
        
        # 3. 작은따옴표를 큰따옴표로 (JSON 키값에만)
        json_text = re.sub(r"'([^']*)':", r'"\1":', json_text)
        
        logger.info(f"[clean_json] JSON 정리 완료 - 길이: {len(json_text)}")
        return json_text
    
    def extract_json_from_text(self, text: str) -> Optional[Dict]:
        """텍스트에서 JSON을 추출하고 파싱하는 함수"""
        logger.info("[extract_json] JSON 추출 시작")
        
        # 코드블록이 포함된 경우 처리
        if '```json' in text.lower():
            # 코드블록 내용만 추출
            match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL | re.IGNORECASE)
            if match:
                text = match.group(1)
                logger.info("[extract_json] 코드블록에서 JSON 추출")
        
        # 1. 정리된 JSON 시도
        cleaned_json = self.clean_json_response(text)
        
        try:
            result = json.loads(cleaned_json)
            logger.info("[extract_json] JSON 파싱 성공")
            return result
        except json.JSONDecodeError as e:
            logger.warning(f"[extract_json] 첫 번째 파싱 실패: {str(e)}")
            logger.debug(f"[extract_json] 문제 위치: {cleaned_json[max(0, e.pos-50):e.pos+50]}")
        
        # 2. 불완전한 JSON 복구 시도
        try:
            # JSON이 중간에 잘린 경우 닫기
            open_braces = cleaned_json.count('{') - cleaned_json.count('}')
            open_brackets = cleaned_json.count('[') - cleaned_json.count(']')
            
            if open_braces > 0 or open_brackets > 0:
                logger.info(f"[extract_json] 불완전한 JSON 감지 - 중괄호: {open_braces}, 대괄호: {open_brackets}")
                
                # 닫는 괄호 추가
                fixed_json = cleaned_json
                fixed_json += ']' * open_brackets + '}' * open_braces
                
                # 마지막 쉼표 제거
                fixed_json = re.sub(r',\s*]', ']', fixed_json)
                fixed_json = re.sub(r',\s*}', '}', fixed_json)
                
                result = json.loads(fixed_json)
                logger.info("[extract_json] 불완전한 JSON 복구 성공")
                return result
                
        except Exception as e:
            logger.warning(f"[extract_json] JSON 복구 실패: {str(e)}")
        
        logger.error("[extract_json] JSON 추출 완전 실패")
        return None
    
    def generate_course_structure_stream(self, 
                                       user_prompt: str,
                                       files: List = None,
                                       settings: Dict[str, Any] = None) -> Generator:
        """파일 텍스트 추출 후 스트리밍 방식으로 코스 구조 생성"""
        try:
            logger.info("[generate_stream] 코스 구조 생성 시작")
            logger.info(f"[generate_stream] 사용자 프롬프트: {user_prompt[:100]}...")
            logger.info(f"[generate_stream] 첨부 파일 수: {len(files) if files else 0}")
            
            # 파일 텍스트 추출
            extracted_text = ""
            if files:
                yield {
                    "type": "progress",
                    "data": f"{len(files)}개 파일 텍스트 추출 중...",
                    "timestamp": datetime.now().isoformat()
                }
                
                extracted_text = self.process_files_to_text(files)
                
                yield {
                    "type": "progress",
                    "data": f"파일 텍스트 추출 완료 ({len(extracted_text)} 글자)",
                    "timestamp": datetime.now().isoformat()
                }
            
            # 설정값 추출
            settings = settings or {}
            logger.info(f"[generate_stream] 설정값: {settings}")
            
            # 프롬프트 생성 - 더 간결하게
            prompt = self.create_concise_course_prompt(
                user_prompt=user_prompt,
                extracted_text=extracted_text,
                education_level=settings.get('education_level', 'middle'),
                difficulty=settings.get('course_difficulty', '중'),
                include_assessment=settings.get('include_assessment', True),
                include_activities=settings.get('include_activities', True),
                include_resources=settings.get('include_resources', False)
            )
            
            logger.info("[generate_stream] Gemini API 호출 시작")
            
            # 스트리밍 응답 생성 - 토큰 제한 증가
            response = self.model.generate_content(
                prompt,
                stream=True,
                generation_config={
                    "temperature": 0.7,
                    "top_p": 0.95,
                    "max_output_tokens": 16384,  # 8192에서 16384로 증가
                },
                safety_settings={
                    "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                    "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                    "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                    "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"
                }
            )
            
            logger.info("[generate_stream] Gemini API 응답 수신 시작")
            
            # 스트리밍 응답 처리
            full_text = ""
            chunk_count = 0
            last_progress_update = 0
            
            for chunk in response:
                chunk_count += 1
                try:
                    # 텍스트 추출 로직
                    chunk_text = ""
                    
                    if hasattr(chunk, 'candidates') and chunk.candidates:
                        candidate = chunk.candidates[0]
                        
                        if hasattr(candidate, 'finish_reason'):
                            if candidate.finish_reason == 2:  # SAFETY
                                logger.warning(f"[generate_stream] 청크 {chunk_count} - 안전성 필터")
                                continue
                            elif candidate.finish_reason == 1:  # STOP (정상 종료)
                                logger.info(f"[generate_stream] 정상 종료 - finish_reason: STOP")
                            elif candidate.finish_reason in [3, 4, 5]:  # OTHER 이유
                                logger.warning(f"[generate_stream] finish_reason: {candidate.finish_reason}")
                        
                        if hasattr(candidate, 'content') and candidate.content and hasattr(candidate.content, 'parts'):
                            for part in candidate.content.parts:
                                if hasattr(part, 'text') and part.text:
                                    chunk_text = part.text
                    elif hasattr(chunk, 'text'):
                        try:
                            chunk_text = chunk.text
                        except ValueError:
                            continue
                    
                    if chunk_text:
                        full_text += chunk_text
                        
                        # 5청크마다 진행 상황 업데이트
                        if chunk_count - last_progress_update >= 5:
                            last_progress_update = chunk_count
                            
                            # JSON 구조 진행 상황 체크
                            if '"chapters"' in full_text:
                                chapter_count = full_text.count('"chapter_title"')
                                yield {
                                    "type": "progress",
                                    "data": f"코스 구조 생성 중... ({chapter_count}개 대단원 생성됨)",
                                    "timestamp": datetime.now().isoformat()
                                }
                            else:
                                yield {
                                    "type": "progress",
                                    "data": f"코스 구조 생성 중... ({chunk_count}번째 청크)",
                                    "timestamp": datetime.now().isoformat()
                                }
                            
                except Exception as e:
                    logger.warning(f"[generate_stream] 청크 {chunk_count} 처리 오류: {str(e)}")
                    continue
            
            logger.info(f"[generate_stream] 스트리밍 완료 - 총 청크: {chunk_count}, 텍스트 길이: {len(full_text)}")
            
            # 응답이 없는 경우
            if not full_text:
                logger.warning("[generate_stream] 응답 텍스트가 비어있음")
                default_structure = self.create_default_structure(user_prompt)
                
                yield {
                    "type": "complete",
                    "data": {
                        "success": True,
                        "structure": default_structure,
                        "statistics": self.calculate_statistics_v2(default_structure),
                        "model_used": self.model_type,
                        "warning": "AI 응답이 비어있어 기본 구조를 생성했습니다."
                    },
                    "timestamp": datetime.now().isoformat()
                }
                return
            
            # JSON 파싱
            course_structure = self.extract_json_from_text(full_text)
            
            if course_structure:
                logger.info(f"[generate_stream] 코스 구조 파싱 성공")
                
                # 구조 검증 및 수정
                course_structure = self.validate_and_fix_structure(course_structure)
                
                # 통계 계산
                stats = self.calculate_statistics_v2(course_structure)
                logger.info(f"[generate_stream] 통계 계산 완료: {stats}")
                
                yield {
                    "type": "complete",
                    "data": {
                        "success": True,
                        "structure": course_structure,
                        "statistics": stats,
                        "model_used": self.model_type
                    },
                    "timestamp": datetime.now().isoformat()
                }
                
            else:
                # JSON 파싱 실패 시 부분 구조라도 생성 시도
                logger.error("[generate_stream] JSON 파싱 실패 - 부분 구조 생성 시도")
                
                partial_structure = self.create_partial_structure_from_text(full_text, user_prompt)
                
                yield {
                    "type": "complete",
                    "data": {
                        "success": True,
                        "structure": partial_structure,
                        "statistics": self.calculate_statistics_v2(partial_structure),
                        "model_used": self.model_type,
                        "warning": "응답이 중간에 잘려서 부분적인 구조만 생성되었습니다."
                    },
                    "timestamp": datetime.now().isoformat()
                }
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            logger.error(f"[generate_stream] 오류 발생: {str(e)}")
            logger.error(f"[generate_stream] 상세 오류:\n{error_trace}")
            
            yield {
                "type": "error",
                "data": {
                    "success": False,
                    "error": f"코스 구조 생성 오류: {str(e)}",
                    "traceback": error_trace
                },
                "timestamp": datetime.now().isoformat()
            }
    
    def validate_and_fix_structure(self, structure: Dict[str, Any]) -> Dict[str, Any]:
        """구조 검증 및 수정 (모델 구조에 맞게)"""
        logger.info("[validate_structure] 구조 검증 시작")
        
        # course 검증
        if 'course' not in structure:
            structure['course'] = {
                "subject_name": "새 코스",
                "target": "중학교 1학년",
                "description": "코스 설명"
            }
        else:
            # 필수 필드 확인
            course = structure['course']
            if 'subject_name' not in course:
                course['subject_name'] = "새 코스"
            if 'target' not in course:
                course['target'] = "대상 미정"
            if 'description' not in course:
                course['description'] = ""
        
        # chapters 검증
        if 'chapters' not in structure or not isinstance(structure['chapters'], list):
            structure['chapters'] = []
        
        # 각 챕터 검증
        for i, chapter in enumerate(structure['chapters']):
            if 'chapter_order' not in chapter:
                chapter['chapter_order'] = i + 1
            if 'chapter_title' not in chapter:
                chapter['chapter_title'] = f"대단원 {i + 1}"
            if 'description' not in chapter:
                chapter['description'] = ""
            if 'subchapters' not in chapter or not isinstance(chapter['subchapters'], list):
                chapter['subchapters'] = []
            
            # 소단원 검증
            for j, subchapter in enumerate(chapter['subchapters']):
                if 'sub_chapter_order' not in subchapter:
                    subchapter['sub_chapter_order'] = j + 1
                if 'sub_chapter_title' not in subchapter:
                    subchapter['sub_chapter_title'] = f"소단원 {j + 1}"
                if 'description' not in subchapter:
                    subchapter['description'] = ""
                if 'chasis' not in subchapter or not isinstance(subchapter['chasis'], list):
                    subchapter['chasis'] = []
                
                # 차시 검증
                for k, chasi in enumerate(subchapter['chasis']):
                    if 'chasi_order' not in chasi:
                        chasi['chasi_order'] = k + 1
                    if 'chasi_title' not in chasi:
                        chasi['chasi_title'] = f"{k + 1}차시"
                    if 'description' not in chasi:
                        chasi['description'] = ""
                    if 'learning_objectives' not in chasi:
                        chasi['learning_objectives'] = ""
                    if 'duration_minutes' not in chasi:
                        chasi['duration_minutes'] = 45
        
        logger.info("[validate_structure] 구조 검증 완료")
        return structure
    
    def create_default_structure(self, user_prompt: str) -> Dict[str, Any]:
        """기본 구조 생성 (모델 구조에 맞게)"""
        logger.info("[create_default] 기본 구조 생성")
        
        return {
            "course": {
                "subject_name": "새 교육 코스",
                "target": "중학교",
                "description": f"사용자 요청: {user_prompt[:100]}..."
            },
            "chapters": [
                {
                    "chapter_title": "1. 기초 개념",
                    "chapter_order": 1,
                    "description": "기본 개념 학습",
                    "subchapters": [
                        {
                            "sub_chapter_title": "1-1. 개념 소개",
                            "sub_chapter_order": 1,
                            "description": "기본 개념 소개",
                            "chasis": [
                                {
                                    "chasi_title": "1차시: 기초 학습",
                                    "chasi_order": 1,
                                    "description": "기초 개념 학습",
                                    "learning_objectives": "- 기본 개념 이해\n- 핵심 내용 파악",
                                    "duration_minutes": 45,
                                    "include_assessment": True,
                                    "include_activities": True,
                                    "include_resources": False
                                }
                            ]
                        }
                    ]
                }
            ]
        }
    
    def create_partial_structure_from_text(self, text: str, user_prompt: str) -> Dict[str, Any]:
        """불완전한 텍스트에서 부분 구조 생성"""
        logger.info("[create_partial] 부분 구조 생성 시도")
        
        structure = {
            "course": {
                "subject_name": "새 교육 코스",
                "target": "대상 미정",
                "description": f"사용자 요청: {user_prompt[:100]}..."
            },
            "chapters": []
        }
        
        # 텍스트에서 코스 정보 추출 시도
        course_match = re.search(r'"course"\s*:\s*\{([^}]+)\}', text)
        if course_match:
            try:
                course_text = "{" + course_match.group(1) + "}"
                course_data = json.loads(course_text)
                structure["course"].update(course_data)
            except:
                pass
        
        # 대단원 추출 시도
        chapter_pattern = r'"chapter_title"\s*:\s*"([^"]+)"'
        chapters = re.findall(chapter_pattern, text)
        
        for i, chapter_title in enumerate(chapters[:3]):  # 최대 3개
            chapter = {
                "chapter_title": chapter_title,
                "chapter_order": i + 1,
                "description": "",
                "subchapters": []
            }
            
            # 기본 소단원 추가
            for j in range(2):  # 각 대단원당 2개 소단원
                subchapter = {
                    "sub_chapter_title": f"{i+1}-{j+1}. 소단원",
                    "sub_chapter_order": j + 1,
                    "description": "",
                    "chasis": []
                }
                
                # 기본 차시 추가
                for k in range(3):  # 각 소단원당 3개 차시
                    chasi = {
                        "chasi_title": f"{k+1}차시",
                        "chasi_order": k + 1,
                        "description": "",
                        "learning_objectives": "학습 목표",
                        "duration_minutes": 45
                    }
                    subchapter["chasis"].append(chasi)
                
                chapter["subchapters"].append(subchapter)
            
            structure["chapters"].append(chapter)
        
        logger.info(f"[create_partial] 부분 구조 생성 완료 - 대단원: {len(structure['chapters'])}개")
        return structure
    
    def calculate_statistics_v2(self, structure: Dict[str, Any]) -> Dict[str, Any]:
        """모델 구조에 맞는 통계 계산"""
        logger.info("[calculate_statistics] 통계 계산 시작")
        
        stats = {
            "total_chapters": 0,
            "total_subchapters": 0,
            "total_chasis": 0,
            "total_duration_hours": 0,
            "total_duration_days": 0,
            "average_chasis_per_subchapter": 0,
            "has_assessment": False,
            "has_activities": False,
            "has_resources": False,
            "complexity_score": 0
        }
        
        if "chapters" in structure:
            stats["total_chapters"] = len(structure["chapters"])
            
            for chapter in structure["chapters"]:
                if "subchapters" in chapter:
                    stats["total_subchapters"] += len(chapter["subchapters"])
                    
                    for subchapter in chapter["subchapters"]:
                        if "chasis" in subchapter:
                            stats["total_chasis"] += len(subchapter["chasis"])
                            
                            for chasi in subchapter["chasis"]:
                                # 시간 계산
                                duration = chasi.get("duration_minutes", 45)
                                stats["total_duration_hours"] += duration / 60
                                
                                # 추가 요소 확인
                                if chasi.get("include_assessment", False):
                                    stats["has_assessment"] = True
                                if chasi.get("include_activities", False):
                                    stats["has_activities"] = True
                                if chasi.get("include_resources", False):
                                    stats["has_resources"] = True
        
        # 통계 계산
        stats["total_duration_hours"] = round(stats["total_duration_hours"], 1)
        stats["total_duration_days"] = round(stats["total_duration_hours"] / 6, 1)  # 하루 6시간 기준
        
        if stats["total_subchapters"] > 0:
            stats["average_chasis_per_subchapter"] = round(
                stats["total_chasis"] / stats["total_subchapters"], 1
            )
        
        # 복잡도 점수 계산 (0-100)
        stats["complexity_score"] = min(100, 
            (stats["total_chapters"] * 10) + 
            (stats["total_subchapters"] * 3) + 
            (stats["total_chasis"])
        )
        
        logger.info(f"[calculate_statistics] 통계 계산 완료: {stats}")
        return stats


@csrf_exempt
@require_http_methods(["POST"])
def generate_course_structure_stream_07211825(self, 
                                   user_prompt: str,
                                   files: List = None,
                                   settings: Dict[str, Any] = None) -> Generator:
    """파일 텍스트 추출 후 스트리밍 방식으로 코스 구조 생성"""
    try:
        logger.info("[generate_stream] 코스 구조 생성 시작")
        logger.info(f"[generate_stream] 사용자 프롬프트: {user_prompt[:100]}...")
        logger.info(f"[generate_stream] 첨부 파일 수: {len(files) if files else 0}")
        
        # 파일 텍스트 추출
        extracted_text = ""
        if files:
            yield {
                "type": "progress",
                "data": f"{len(files)}개 파일 텍스트 추출 중...",
                "timestamp": datetime.now().isoformat()
            }
            
            extracted_text = self.process_files_to_text(files)
            
            yield {
                "type": "progress",
                "data": f"파일 텍스트 추출 완료 ({len(extracted_text)} 글자)",
                "timestamp": datetime.now().isoformat()
            }
        
        # 설정값 추출
        settings = settings or {}
        logger.info(f"[generate_stream] 설정값: {settings}")
        
        # 전체 코스 생성 여부 판단
        is_full_course_mode = False
        full_course_keywords = ['전체', '모든', '완전한', '파일을 바탕으로', '첨부한 파일', 'pdf', 'PDF']
        
        if extracted_text and any(keyword in user_prompt for keyword in full_course_keywords):
            is_full_course_mode = True
            logger.info("[generate_stream] 전체 코스 생성 모드 활성화")
        
        # 프롬프트 생성
        if is_full_course_mode:
            # 전체 코스 생성 프롬프트
            prompt = self.create_full_course_prompt(
                user_prompt=user_prompt,
                extracted_text=extracted_text,
                education_level=settings.get('education_level', 'middle'),
                difficulty=settings.get('course_difficulty', '중'),
                include_assessment=settings.get('include_assessment', True),
                include_activities=settings.get('include_activities', True),
                include_resources=settings.get('include_resources', False)
            )
            # 토큰 제한을 더 크게 설정
            max_tokens = 32768  # 전체 코스 생성 시 더 많은 토큰 허용
        else:
            # 기존 간결한 프롬프트 사용
            prompt = self.create_concise_course_prompt(
                user_prompt=user_prompt,
                extracted_text=extracted_text,
                education_level=settings.get('education_level', 'middle'),
                difficulty=settings.get('course_difficulty', '중'),
                include_assessment=settings.get('include_assessment', True),
                include_activities=settings.get('include_activities', True),
                include_resources=settings.get('include_resources', False)
            )
            max_tokens = 16384
        
        logger.info(f"[generate_stream] 프롬프트 길이: {len(prompt)}, 최대 토큰: {max_tokens}")
        logger.info("[generate_stream] Gemini API 호출 시작")
        
        # 스트리밍 응답 생성
        response = self.model.generate_content(
            prompt,
            stream=True,
            generation_config={
                "temperature": 0.7 if not is_full_course_mode else 0.5,  # 전체 코스 모드일 때는 더 정확하게
                "top_p": 0.95 if not is_full_course_mode else 0.9,
                "max_output_tokens": max_tokens,
            },
            safety_settings={
                "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"
            }
        )
        
        logger.info("[generate_stream] Gemini API 응답 수신 시작")
        
        # 스트리밍 응답 처리
        full_text = ""
        chunk_count = 0
        last_progress_update = 0
        chapters_count = 0
        
        for chunk in response:
            chunk_count += 1
            try:
                # 텍스트 추출 로직
                chunk_text = ""
                
                if hasattr(chunk, 'candidates') and chunk.candidates:
                    candidate = chunk.candidates[0]
                    
                    if hasattr(candidate, 'finish_reason'):
                        if candidate.finish_reason == 2:  # SAFETY
                            logger.warning(f"[generate_stream] 청크 {chunk_count} - 안전성 필터")
                            continue
                        elif candidate.finish_reason == 1:  # STOP (정상 종료)
                            logger.info(f"[generate_stream] 정상 종료 - finish_reason: STOP")
                        elif candidate.finish_reason in [3, 4, 5]:  # OTHER 이유
                            logger.warning(f"[generate_stream] finish_reason: {candidate.finish_reason}")
                            # 토큰 제한에 도달했을 가능성이 있으므로 현재까지의 결과로 진행
                            if candidate.finish_reason == 5 and full_text:  # LENGTH
                                logger.warning("[generate_stream] 토큰 제한 도달, 현재까지의 결과로 처리")
                                break
                    
                    if hasattr(candidate, 'content') and candidate.content and hasattr(candidate.content, 'parts'):
                        for part in candidate.content.parts:
                            if hasattr(part, 'text') and part.text:
                                chunk_text = part.text
                elif hasattr(chunk, 'text'):
                    try:
                        chunk_text = chunk.text
                    except ValueError:
                        continue
                
                if chunk_text:
                    full_text += chunk_text
                    
                    # 진행 상황 업데이트 로직 개선
                    current_chapters = full_text.count('"chapter_title"')
                    
                    # 5청크마다 또는 챕터 수가 변경되었을 때 업데이트
                    if (chunk_count - last_progress_update >= 5) or (current_chapters > chapters_count):
                        last_progress_update = chunk_count
                        chapters_count = current_chapters
                        
                        # JSON 구조 진행 상황 체크
                        if '"chapters"' in full_text:
                            # 더 상세한 진행 정보 제공
                            subchapter_count = full_text.count('"sub_chapter_title"')
                            chasi_count = full_text.count('"chasi_title"')
                            
                            if is_full_course_mode:
                                progress_msg = f"전체 코스 생성 중... ({chapters_count}개 대단원, {subchapter_count}개 소단원, {chasi_count}개 차시)"
                            else:
                                progress_msg = f"코스 구조 생성 중... ({chapters_count}개 대단원 생성됨)"
                            
                            yield {
                                "type": "progress",
                                "data": progress_msg,
                                "timestamp": datetime.now().isoformat()
                            }
                        else:
                            yield {
                                "type": "progress",
                                "data": f"코스 구조 생성 중... ({chunk_count}번째 청크)",
                                "timestamp": datetime.now().isoformat()
                            }
                        
            except Exception as e:
                logger.warning(f"[generate_stream] 청크 {chunk_count} 처리 오류: {str(e)}")
                continue
        
        logger.info(f"[generate_stream] 스트리밍 완료 - 총 청크: {chunk_count}, 텍스트 길이: {len(full_text)}")
        
        # 응답이 없는 경우
        if not full_text:
            logger.warning("[generate_stream] 응답 텍스트가 비어있음")
            default_structure = self.create_default_structure(user_prompt)
            
            yield {
                "type": "complete",
                "data": {
                    "success": True,
                    "structure": default_structure,
                    "statistics": self.calculate_statistics_v2(default_structure),
                    "model_used": self.model_type,
                    "warning": "AI 응답이 비어있어 기본 구조를 생성했습니다."
                },
                "timestamp": datetime.now().isoformat()
            }
            return
        
        # JSON 파싱
        course_structure = self.extract_json_from_text(full_text)
        
        if course_structure:
            logger.info(f"[generate_stream] 코스 구조 파싱 성공")
            
            # 구조 검증 및 수정
            course_structure = self.validate_and_fix_structure(course_structure)
            
            # 통계 계산
            stats = self.calculate_statistics_v2(course_structure)
            logger.info(f"[generate_stream] 통계 계산 완료: {stats}")
            
            # 전체 코스 모드에서 예상 차시 수와 실제 차시 수 비교
            if is_full_course_mode and extracted_text:
                # PDF에서 추출한 차시 수 계산 (간단한 휴리스틱)
                expected_chasi_count = extracted_text.count('차시')
                actual_chasi_count = stats['total_chasis']
                
                if actual_chasi_count < expected_chasi_count * 0.8:  # 80% 미만이면 경고
                    logger.warning(f"[generate_stream] 차시 누락 가능성 - 예상: {expected_chasi_count}, 실제: {actual_chasi_count}")
                    stats['warning'] = f"일부 차시가 누락되었을 수 있습니다. (생성된 차시: {actual_chasi_count}개)"
            
            yield {
                "type": "complete",
                "data": {
                    "success": True,
                    "structure": course_structure,
                    "statistics": stats,
                    "model_used": self.model_type,
                    "mode": "full" if is_full_course_mode else "concise"
                },
                "timestamp": datetime.now().isoformat()
            }
            
        else:
            # JSON 파싱 실패 시 부분 구조라도 생성 시도
            logger.error("[generate_stream] JSON 파싱 실패 - 부분 구조 생성 시도")
            logger.debug(f"[generate_stream] 파싱 실패한 텍스트 일부: {full_text[:1000]}...")
            
            partial_structure = self.create_partial_structure_from_text(full_text, user_prompt)
            
            yield {
                "type": "complete",
                "data": {
                    "success": True,
                    "structure": partial_structure,
                    "statistics": self.calculate_statistics_v2(partial_structure),
                    "model_used": self.model_type,
                    "warning": "응답이 중간에 잘려서 부분적인 구조만 생성되었습니다. 다시 시도해보세요.",
                    "mode": "partial"
                },
                "timestamp": datetime.now().isoformat()
            }
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"[generate_stream] 오류 발생: {str(e)}")
        logger.error(f"[generate_stream] 상세 오류:\n{error_trace}")
        
        yield {
            "type": "error",
            "data": {
                "success": False,
                "error": f"코스 구조 생성 오류: {str(e)}",
                "traceback": error_trace
            },
            "timestamp": datetime.now().isoformat()
        }


def design_course_structure_stream(request):
    """스트리밍 방식의 코스 구조 설계 API"""
    logger.info("="*50)
    logger.info("[API] 코스 구조 설계 요청 수신 (스트리밍)")
    logger.info(f"[API] 요청 메서드: {request.method}")
    logger.info(f"[API] Content-Type: {request.content_type}")
    
    try:
        # 모델 타입 선택
        model_type = request.POST.get('model_type', 'flash')
        logger.info(f"[API] 모델 타입: {model_type}")
        
        designer = CourseDesigner(model_type=model_type)
        
        # 프롬프트 가져오기
        user_prompt = request.POST.get('prompt', '')
        logger.info(f"[API] 프롬프트 길이: {len(user_prompt)}")
        logger.info(f"[API] 프롬프트 내용: {user_prompt[:200]}...")
        
        if not user_prompt:
            logger.error("[API] 프롬프트가 비어있음")
            return JsonResponse({
                'success': False,
                'error': '프롬프트를 입력해주세요.'
            })
        
        # 고급 설정값 가져오기
        settings = {
            'education_level': request.POST.get('education_level', 'middle'),
            'course_difficulty': request.POST.get('course_difficulty', '중'),
            'include_assessment': request.POST.get('include_assessment', 'true').lower() == 'true',
            'include_activities': request.POST.get('include_activities', 'true').lower() == 'true',
            'include_resources': request.POST.get('include_resources', 'false').lower() == 'true'
        }
        logger.info(f"[API] 설정값: {settings}")
        
        # 파일 처리
        files = request.FILES.getlist('files')
        
        if files:
            logger.info(f"[API] 업로드된 파일 수: {len(files)}")
            for f in files:
                logger.info(f"[API] 파일: {f.name} ({f.size} bytes, {f.content_type})")
        
        # 스트리밍 응답 생성
        logger.info("[API] 스트리밍 응답 생성 시작")
        
        def generate():
            try:
                logger.info("[API-Stream] 제너레이터 시작")
                message_count = 0
                
                for chunk in designer.generate_course_structure_stream(
                    user_prompt=user_prompt,
                    files=files,
                    settings=settings
                ):
                    message_count += 1
                    chunk_data = f"data: {json.dumps(chunk, ensure_ascii=False)}\n\n"
                    logger.debug(f"[API-Stream] 메시지 {message_count} 전송 - 타입: {chunk.get('type')}")
                    yield chunk_data
                
                logger.info(f"[API-Stream] 스트리밍 완료 - 총 메시지 수: {message_count}")
                
            except Exception as e:
                logger.error(f"[API-Stream] 스트리밍 중 오류: {str(e)}", exc_info=True)
                error_data = {
                    "type": "error",
                    "data": {
                        "success": False,
                        "error": str(e)
                    },
                    "timestamp": datetime.now().isoformat()
                }
                yield f"data: {json.dumps(error_data, ensure_ascii=False)}\n\n"
        
        response = StreamingHttpResponse(
            generate(),
            content_type='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',
            }
        )
        
        logger.info("[API] StreamingHttpResponse 반환")
        return response
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"[API] 전체 오류: {str(e)}")
        logger.error(f"[API] 상세 오류:\n{error_trace}")
        
        return JsonResponse({
            'success': False,
            'error': f'서버 오류: {str(e)}',
            'detail': error_trace
        })


# 비스트리밍 버전 (폴백용)
@csrf_exempt
@require_http_methods(["POST"])
def design_course_structure(request):
    """기존 방식의 코스 구조 설계 API (비스트리밍)"""
    logger.info("[API-Regular] 비스트리밍 API 호출")
    
    try:
        # 모델 타입 선택
        model_type = request.POST.get('model_type', 'flash')
        designer = CourseDesigner(model_type=model_type)
        
        # 프롬프트 가져오기
        user_prompt = request.POST.get('prompt', '')
        
        if not user_prompt:
            return JsonResponse({
                'success': False,
                'error': '프롬프트를 입력해주세요.'
            })
        
        # 설정값
        settings = {
            'education_level': request.POST.get('education_level', 'middle'),
            'difficulty': request.POST.get('course_difficulty', '중'),
            'include_assessment': request.POST.get('include_assessment', 'true').lower() == 'true',
            'include_activities': request.POST.get('include_activities', 'true').lower() == 'true',
            'include_resources': request.POST.get('include_resources', 'false').lower() == 'true'
        }
        
        # 파일 텍스트 추출
        extracted_text = ""
        files = request.FILES.getlist('files')
        if files:
            extracted_text = designer.process_files_to_text(files)
        
        # 프롬프트 생성
        prompt = designer.create_concise_course_prompt(
            user_prompt=user_prompt,
            extracted_text=extracted_text,
            **settings
        )
        
        # 비스트리밍 응답 생성
        response = designer.model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.7,
                "top_p": 0.95,
                "max_output_tokens": 16384,
            },
            safety_settings={
                "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"
            }
        )
        
        # 응답 텍스트 추출
        full_text = ""
        try:
            if hasattr(response, 'candidates') and response.candidates:
                for candidate in response.candidates:
                    if hasattr(candidate, 'content') and candidate.content:
                        if hasattr(candidate.content, 'parts'):
                            for part in candidate.content.parts:
                                if hasattr(part, 'text') and part.text:
                                    full_text += part.text
            
            if not full_text and hasattr(response, 'text'):
                try:
                    full_text = response.text
                except ValueError:
                    pass
                    
        except Exception as e:
            logger.error(f"[API-Regular] 응답 추출 오류: {str(e)}")
        
        if not full_text:
            # 기본 구조 반환
            default_structure = designer.create_default_structure(user_prompt)
            
            return JsonResponse({
                'success': True,
                'structure': default_structure,
                'statistics': designer.calculate_statistics_v2(default_structure),
                'model_used': model_type,
                'warning': '안전성 필터로 인해 AI 응답을 받을 수 없어 기본 구조를 생성했습니다.'
            })
        
        # JSON 파싱
        course_structure = designer.extract_json_from_text(full_text)
        
        if course_structure:
            course_structure = designer.validate_and_fix_structure(course_structure)
            stats = designer.calculate_statistics_v2(course_structure)
            
            return JsonResponse({
                'success': True,
                'structure': course_structure,
                'statistics': stats,
                'model_used': model_type
            })
        else:
            # 부분 구조 시도
            partial_structure = designer.create_partial_structure_from_text(full_text, user_prompt)
            
            return JsonResponse({
                'success': True,
                'structure': partial_structure,
                'statistics': designer.calculate_statistics_v2(partial_structure),
                'model_used': model_type,
                'warning': 'AI 응답 파싱에 실패하여 부분 구조를 생성했습니다.',
                'raw_response': full_text[:500] if full_text else "응답 없음"
            })
            
    except Exception as e:
        logger.error(f"[API-Regular] 오류: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'서버 오류: {str(e)}'
        })